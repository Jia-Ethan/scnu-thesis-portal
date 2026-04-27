from pathlib import Path

from docx import Document

from backend.app.contracts import CapabilityFlags
from backend.app.services.parse import normalize_text_input, parse_docx_file
from backend.app.services.precheck import run_precheck


FIXTURE = Path(__file__).resolve().parent / "fixtures" / "sample-thesis.docx"
MISSING_FIXTURE = Path(__file__).resolve().parent / "fixtures" / "missing-sections.docx"


def capabilities():
    return CapabilityFlags(docx_export=True, profile="undergraduate")


def test_parse_docx_extracts_cover_and_expected_sections():
    thesis = parse_docx_file(FIXTURE, capabilities())

    assert thesis.cover.title == "基于结构化映射的本科论文生成示例"
    assert thesis.abstract_cn.content
    assert thesis.abstract_en.content
    assert thesis.references
    assert thesis.appendices
    assert thesis.acknowledgements
    assert thesis.body_sections[0].title == "引言"


def test_parse_docx_tracks_missing_sections_without_fabrication():
    thesis = parse_docx_file(MISSING_FIXTURE, capabilities())

    assert "appendices" in thesis.missing_sections
    assert "acknowledgements" in thesis.missing_sections
    assert "cover.advisor" in thesis.missing_sections
    assert thesis.references[0].normalized_text == "示例作者. 规范化导出实践[J]."

    precheck = run_precheck(thesis)
    assert precheck.summary.can_confirm is True
    assert any(issue.code == "APPENDICES_BLANK" for issue in precheck.issues)
    assert any(issue.code == "ACKNOWLEDGEMENTS_BLANK" for issue in precheck.issues)


def test_text_normalize_keeps_missing_sections_as_warnings():
    thesis = normalize_text_input("# 引言\n\n正文内容。" * 20, capabilities())

    assert thesis.body_sections
    assert "abstract_cn" in thesis.missing_sections
    assert "references" in thesis.missing_sections
    assert "未识别到中文摘要，导出时会保留摘要章节留白。" in thesis.warnings

    precheck = run_precheck(thesis)
    assert precheck.summary.can_confirm is True
    assert precheck.summary.blocking_count == 0
    assert any(issue.code == "ABSTRACT_CN_BLANK" for issue in precheck.issues)


def test_docx_and_text_inputs_can_converge_to_same_semantics():
    docx_thesis = parse_docx_file(MISSING_FIXTURE, capabilities())
    text_thesis = normalize_text_input(
        "\n".join(
            [
                "论文题目：面向规范映射的本科论文导出基线",
                "学生姓名：张三",
                "学号：2020123456",
                "学院：计算机学院",
                "专业：网络工程",
                "摘要",
                "本文用于验证缺失章节时仍能保留结构留白，不自动补写不存在的内容。",
                "关键词：导出规范，章节映射，留白策略",
                "第一章 引言",
                "这里是引言正文，用于验证正文仍可被识别并导出。",
                "1.1 研究背景",
                "这里是研究背景，用于验证多级标题目录联动。",
                "参考文献",
                "【1】示例作者. 规范化导出实践[J].",
            ]
        ),
        capabilities(),
    )

    assert docx_thesis.cover.title == text_thesis.cover.title
    assert [section.title for section in docx_thesis.body_sections] == [section.title for section in text_thesis.body_sections]
    assert [item.normalized_text for item in docx_thesis.references] == [item.normalized_text for item in text_thesis.references]
    assert docx_thesis.missing_sections == text_thesis.missing_sections


def test_scnu_structure_rules_flag_real_world_front_matter_and_body_risks(tmp_path):
    path = tmp_path / "problematic-thesis.docx"
    document = Document()
    for line in [
        "华南师范大学",
        "本科毕业论文",
        "论文题目：叠滘龙舟非遗文旅消费群体画像与业态升级研究 作者：张三",
        "指导教师：________________",
        "学生姓名：________________",
        "学号：________________",
        "学院：__________________",
        "专业：__________________",
        "班级：__________________",
        "毕业时间：________________",
    ]:
        document.add_paragraph(line)
    document.add_heading("摘  要", level=1)
    document.add_paragraph("关键词：__________________")
    document.add_heading("Abstract", level=1)
    document.add_paragraph("Keywords：__________________")
    document.add_heading("目  录", level=1)
    document.add_heading("1 宏观趋势：Z世代崛起重塑文旅消费", level=1)
    document.add_paragraph("【摘要】真正的中文摘要被放进正文开头。" + "这是一个异常长段落。" * 90 + "[2] 图1 展示消费画像，表1 展示样本结构。")
    document.add_heading("2 政策推进“旅游+”和“+旅游”", level=1)
    document.add_paragraph("一、研究背景\n（一）研究目的\n3.1 层级混用\n有效样本16份，女性占比68.97%，需核对数据口径。")
    document.add_heading("参考文献", level=1)
    document.add_heading("附录", level=1)
    document.add_heading("致谢", level=1)
    document.save(path)

    thesis = parse_docx_file(path, capabilities())
    precheck = run_precheck(thesis)
    codes = {item.code for item in precheck.issues}

    assert "cover.advisor" in thesis.missing_sections
    assert "ABSTRACT_CN_BLANK" in codes
    assert "ABSTRACT_EN_BLANK" in codes
    assert "TOC_EMPTY" in codes
    assert "ABSTRACT_CN_MISPLACED" in codes
    assert "CITATIONS_WITHOUT_REFERENCES" in codes
    assert "HEADING_NUMBERING_MIXED" in codes
    assert "LONG_PARAGRAPHS" in codes
    assert "FIGURE_CAPTION_WITHOUT_OBJECT" in codes
    assert "TABLE_CAPTION_WITHOUT_TABLE" in codes
    assert "APPENDIX_EMPTY" in codes
    assert "ACKNOWLEDGEMENTS_EMPTY" in codes
    assert "DATA_CONSISTENCY_REVIEW" in codes
