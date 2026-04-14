from __future__ import annotations

import argparse
import json
import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any
from zipfile import ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn

from backend.app.services.export import extract_header_title

BANNED_PLACEHOLDERS = [
    "待补充论文题目",
    "未填写",
    "请在 Word 中右键更新目录",
]
WORD_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


@dataclass
class CheckResult:
    id: str
    status: str
    message: str
    details: dict[str, Any]


@dataclass
class ComplianceReport:
    path: str
    results: list[CheckResult]

    @property
    def summary(self) -> dict[str, int]:
        counts = {"PASS": 0, "MANUAL_REVIEW": 0, "NOT_SUPPORTED": 0}
        for result in self.results:
            counts[result.status] = counts.get(result.status, 0) + 1
        return counts

    def to_dict(self) -> dict[str, Any]:
        return {
            "path": self.path,
            "summary": self.summary,
            "results": [asdict(item) for item in self.results],
        }


def read_docx_parts(path: Path) -> dict[str, str]:
    with ZipFile(path) as archive:
        return {
            name: archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.endswith(".xml") or name.endswith(".rels")
        }


def approx(value: float, expected: float, tolerance: float = 0.08) -> bool:
    return abs(value - expected) <= tolerance


def make_result(results: list[CheckResult], id: str, status: str, message: str, **details: Any) -> None:
    results.append(CheckResult(id=id, status=status, message=message, details=details))


def extract_rfonts_from_run(run) -> dict[str, Any]:
    rpr = run._element.rPr
    rfonts = rpr.rFonts if rpr is not None else None
    size = run.font.size.pt if run.font.size is not None else None
    bold = bool(run.bold)
    if rfonts is None:
        return {"ascii": None, "hAnsi": None, "eastAsia": None, "size": size, "bold": bold}
    return {
        "ascii": rfonts.get(qn("w:ascii")),
        "hAnsi": rfonts.get(qn("w:hAnsi")),
        "eastAsia": rfonts.get(qn("w:eastAsia")),
        "size": size,
        "bold": bold,
    }


def extract_first_run_props(xml_bytes: bytes) -> dict[str, Any]:
    root = ET.fromstring(xml_bytes)
    run = root.find(".//w:r", WORD_NS)
    if run is None:
        return {}
    rpr = run.find("w:rPr", WORD_NS)
    if rpr is None:
        return {}
    rfonts = rpr.find("w:rFonts", WORD_NS)
    size = rpr.find("w:sz", WORD_NS)
    return {
        "ascii": rfonts.attrib.get(f"{{{WORD_NS['w']}}}ascii") if rfonts is not None else None,
        "hAnsi": rfonts.attrib.get(f"{{{WORD_NS['w']}}}hAnsi") if rfonts is not None else None,
        "eastAsia": rfonts.attrib.get(f"{{{WORD_NS['w']}}}eastAsia") if rfonts is not None else None,
        "size": (int(size.attrib.get(f"{{{WORD_NS['w']}}}val")) / 2) if size is not None else None,
        "bold": (lambda node: node is not None and node.attrib.get(f"{{{WORD_NS['w']}}}val", "1") not in {"0", "false", "False"})(rpr.find("w:b", WORD_NS)),
    }


def nonempty_paragraphs(document: Document):
    return [(index, paragraph) for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip()]


def find_paragraph_index(paragraphs, *candidates: str) -> int | None:
    normalized_candidates = {normalize_for_search(candidate) for candidate in candidates if candidate}
    for index, paragraph in paragraphs:
        if normalize_for_search(paragraph.text) in normalized_candidates:
            return index
    return None


def find_cover_title(document: Document) -> str:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("论文题目："):
            value = text.split("：", 1)[1].strip()
            if value and not re.fullmatch(r"_+", value):
                return value
            return ""
    return ""


def first_nonempty_after(paragraphs, start_index: int | None):
    if start_index is None:
        return None
    for index, paragraph in paragraphs:
        if index > start_index and paragraph.text.strip():
            return paragraph
    return None


def check_docx(path: Path) -> ComplianceReport:
    document = Document(path)
    parts = read_docx_parts(path)
    document_xml = parts.get("word/document.xml", "")
    paragraphs = nonempty_paragraphs(document)
    results: list[CheckResult] = []

    sections = list(document.sections)
    if not sections:
        make_result(results, "page_size", "MANUAL_REVIEW", "文档未包含 section 信息。")
        make_result(results, "margins_gutter", "MANUAL_REVIEW", "无法检查页边距与装订线。")
    else:
        first_section = sections[0]
        page_size_ok = approx(first_section.page_width.cm, 21) and approx(first_section.page_height.cm, 29.7)
        make_result(
            results,
            "page_size",
            "PASS" if page_size_ok else "MANUAL_REVIEW",
            "页面尺寸已检查。" if page_size_ok else "页面尺寸与 A4 不一致。",
            page_width_cm=round(first_section.page_width.cm, 3),
            page_height_cm=round(first_section.page_height.cm, 3),
        )
        margins_ok = all(
            [
                approx(first_section.top_margin.cm, 2.5),
                approx(first_section.bottom_margin.cm, 2.5),
                approx(first_section.left_margin.cm, 2),
                approx(first_section.right_margin.cm, 2),
                approx(first_section.gutter.cm, 0.5),
            ]
        )
        make_result(
            results,
            "margins_gutter",
            "PASS" if margins_ok else "MANUAL_REVIEW",
            "页边距与装订线已检查。" if margins_ok else "页边距或装订线与规范不一致。",
            top_cm=round(first_section.top_margin.cm, 3),
            bottom_cm=round(first_section.bottom_margin.cm, 3),
            left_cm=round(first_section.left_margin.cm, 3),
            right_cm=round(first_section.right_margin.cm, 3),
            gutter_cm=round(first_section.gutter.cm, 3),
        )

    cn_index = find_paragraph_index(paragraphs, "摘  要", "摘要", "中文摘要")
    cn_heading = next((paragraph for index, paragraph in paragraphs if index == cn_index), None)
    cn_body = first_nonempty_after(paragraphs, cn_index)
    cn_heading_props = extract_rfonts_from_run(cn_heading.runs[0]) if cn_heading and cn_heading.runs else {}
    cn_body_props = extract_rfonts_from_run(cn_body.runs[0]) if cn_body and cn_body.runs else {}
    cn_ok = bool(cn_heading and cn_body) and cn_heading_props.get("eastAsia") == "黑体" and approx(float(cn_heading_props.get("size") or 0), 18, 0.2) and cn_body_props.get("eastAsia") == "宋体" and approx(float(cn_body_props.get("size") or 0), 12, 0.2)
    make_result(results, "cn_abstract_style", "PASS" if cn_ok else "MANUAL_REVIEW", "中文摘要样式已检查。" if cn_ok else "中文摘要标题或正文样式不符合预期。", heading=cn_heading_props, body=cn_body_props)

    en_index = find_paragraph_index(paragraphs, "Abstract")
    en_heading = next((paragraph for index, paragraph in paragraphs if index == en_index), None)
    en_body = first_nonempty_after(paragraphs, en_index)
    en_body_props = extract_rfonts_from_run(en_body.runs[0]) if en_body and en_body.runs else {}
    en_ok = bool(en_heading and en_body) and en_body_props.get("eastAsia") == "Times New Roman" and en_body_props.get("ascii") == "Times New Roman" and approx(float(en_body_props.get("size") or 0), 12, 0.2)
    make_result(results, "en_abstract_style", "PASS" if en_ok else "MANUAL_REVIEW", "英文摘要样式已检查。" if en_ok else "英文摘要正文样式不符合预期。", body=en_body_props)

    toc_ok = 'TOC \\o "1-4" \\h \\z \\u' in document_xml
    make_result(results, "toc_field", "PASS" if toc_ok else "MANUAL_REVIEW", "目录字段已存在。" if toc_ok else "未检测到目录字段。")

    header_files = sorted(name for name in parts if name.startswith("word/header"))
    footer_files = sorted(name for name in parts if name.startswith("word/footer"))
    cover_title = find_cover_title(document)
    expected_header_title = extract_header_title(cover_title)
    header_texts: list[str] = []
    header_props: dict[str, Any] = {}
    for name in header_files:
        xml_text = parts[name]
        text = "".join(ET.fromstring(xml_text).itertext()).strip()
        if not text:
            continue
        header_texts.append(text)
        if not header_props:
            header_props = extract_first_run_props(xml_text.encode("utf-8"))
    header_ok = (bool(expected_header_title) and expected_header_title in header_texts) or (not expected_header_title and not header_texts)
    make_result(results, "header_title", "PASS" if header_ok else "MANUAL_REVIEW", "页眉内容已检查。" if header_ok else "页眉未正确写入论文主标题。", header_texts=header_texts, expected_header_title=expected_header_title)

    footer_has_page = any("PAGE" in parts[name] for name in footer_files)
    make_result(results, "footer_page_field", "PASS" if footer_has_page else "MANUAL_REVIEW", "页脚页码字段已存在。" if footer_has_page else "未检测到页码字段。")

    toc_index = find_paragraph_index(paragraphs, "目  录", "目录")
    ref_index = find_paragraph_index(paragraphs, "参考文献")
    appendix_index = find_paragraph_index(paragraphs, "附录")
    ack_index = find_paragraph_index(paragraphs, "致谢")
    body_index = next(
        (
            index
            for index, paragraph in paragraphs
            if paragraph.style.name.startswith("Heading")
            and normalize_for_search(paragraph.text) not in {normalize_for_search(token) for token in ["摘  要", "摘要", "中文摘要", "Abstract", "目  录", "目录", "参考文献", "附录", "致谢"]}
        ),
        None,
    )
    order_ok = all(index is not None for index in [cn_index, en_index, toc_index, body_index, ref_index])
    if order_ok:
        order_ok = cn_index < en_index < toc_index < body_index < ref_index
        if appendix_index is not None:
            order_ok = order_ok and ref_index < appendix_index
        if ack_index is not None and appendix_index is not None:
            order_ok = order_ok and appendix_index < ack_index
        elif ack_index is not None:
            order_ok = order_ok and ref_index < ack_index
    make_result(results, "section_order", "PASS" if order_ok else "MANUAL_REVIEW", "文档顺序已符合当前口径。" if order_ok else "文档顺序缺项或顺序不符合当前口径。", cn_index=cn_index, en_index=en_index, toc_index=toc_index, body_index=body_index, ref_index=ref_index, appendix_index=appendix_index, ack_index=ack_index)

    cover_field_line = any(paragraph.text.strip().startswith("论文题目：") for _, paragraph in paragraphs)
    advisor_field_line = any(paragraph.text.strip().startswith("指导教师：") for _, paragraph in paragraphs)
    cover_ok = bool(find_paragraph_index(paragraphs, "本科毕业论文")) and cover_field_line and advisor_field_line
    make_result(results, "official_cover_absence", "PASS" if cover_ok else "MANUAL_REVIEW", "正式封面已纳入主线。" if cover_ok else "未检测到正式封面关键字段。")

    make_result(results, "notes_support", "NOT_SUPPORTED", "注释编号与页末注规则仍需人工复核。", reason="当前导出链路仅保留结构位与风险提示。")
    make_result(results, "figure_table_captions", "NOT_SUPPORTED", "图题、表题与复杂对象仍需人工复核。", reason="复杂对象会被标记为人工复核项。")

    return ComplianceReport(path=str(path), results=results)


def normalize_for_search(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def gather_document_text(document_xml: str) -> str:
    return "".join(re.findall(r"<w:t[^>]*>(.*?)</w:t>", document_xml))


def collect_section_formats(document_xml: str) -> list[tuple[str | None, str | None]]:
    formats: list[tuple[str | None, str | None]] = []
    for match in re.finditer(r"<w:sectPr\b.*?</w:sectPr>", document_xml, flags=re.DOTALL):
        sect = match.group(0)
        pg = re.search(r"<w:pgNumType\b([^>]*)/>", sect)
        if not pg:
            formats.append((None, None))
            continue
        attrs = pg.group(1)
        fmt = re.search(r'w:fmt="([^"]+)"', attrs)
        start = re.search(r'w:start="([^"]+)"', attrs)
        formats.append((fmt.group(1) if fmt else None, start.group(1) if start else None))
    return formats


def build_report(path: Path) -> dict[str, object]:
    if not path.exists() or path.suffix.lower() != ".docx":
        return {"status": "NOT_SUPPORTED", "reasons": ["文件不存在或不是 .docx。"]}

    document = Document(path)
    parts = read_docx_parts(path)
    document_xml = parts.get("word/document.xml", "")
    styles_xml = parts.get("word/styles.xml", "")
    settings_xml = parts.get("word/settings.xml", "")
    headers = {name: xml for name, xml in parts.items() if name.startswith("word/header")}
    footers = {name: xml for name, xml in parts.items() if name.startswith("word/footer")}
    paragraphs = nonempty_paragraphs(document)

    text = normalize_for_search(gather_document_text(document_xml))
    reasons: list[str] = []
    manual_review: list[str] = []

    if not document_xml:
        return {"status": "NOT_SUPPORTED", "reasons": ["文档主体 XML 缺失。"]}

    if not all(token in text for token in ["华南师范大学", "本科毕业论文", "论文题目", "指导教师", "学生姓名", "学号", "学院", "专业", "班级"]):
        reasons.append("正式封面字段未完整落位。")

    cn_index = find_paragraph_index(paragraphs, "摘  要", "摘要", "中文摘要")
    en_index = find_paragraph_index(paragraphs, "Abstract")
    toc_index = find_paragraph_index(paragraphs, "目  录", "目录")
    ref_index = find_paragraph_index(paragraphs, "参考文献")
    appendix_index = find_paragraph_index(paragraphs, "附录")
    ack_index = find_paragraph_index(paragraphs, "致谢")
    body_index = next(
        (
            index
            for index, paragraph in paragraphs
            if paragraph.style.name.startswith("Heading")
            and normalize_for_search(paragraph.text) not in {normalize_for_search(token) for token in ["摘  要", "摘要", "中文摘要", "Abstract", "目  录", "目录", "参考文献", "附录", "致谢"]}
        ),
        None,
    )
    order_ok = all(index is not None for index in [cn_index, en_index, toc_index, body_index, ref_index])
    if order_ok:
        order_ok = cn_index < en_index < toc_index < body_index < ref_index
        if appendix_index is not None:
            order_ok = order_ok and ref_index < appendix_index
        if ack_index is not None and appendix_index is not None:
            order_ok = order_ok and appendix_index < ack_index
        elif ack_index is not None:
            order_ok = order_ok and ref_index < ack_index
    if not order_ok:
        reasons.append("摘要 / 目录 / 参考文献 / 附录 / 致谢顺序不符合新主线。")

    if 'TOC\\o"1-4"\\h\\z\\u' not in normalize_for_search(document_xml):
        reasons.append("目录字段缺失。")
    if 'w:updateFieldsw:val="true"' not in normalize_for_search(settings_xml):
        reasons.append("Word 自动更新目录字段设置缺失。")

    section_formats = collect_section_formats(document_xml)
    if not any(fmt == "upperRoman" and start == "1" for fmt, start in section_formats):
        reasons.append("前置部分未检测到大写罗马页码分节。")
    if not any(fmt == "decimal" and start == "1" for fmt, start in section_formats):
        reasons.append("正文未检测到阿拉伯页码重启分节。")

    if not headers:
        reasons.append("页眉文件缺失。")
    if not footers:
        reasons.append("页脚文件缺失。")
    if footers and not any("PAGE" in xml for xml in footers.values()):
        reasons.append("页脚页码字段缺失。")

    normalized_styles = normalize_for_search(styles_xml)
    if "Heading1" not in styles_xml and 'w:namew:val="heading1"' not in normalized_styles:
        reasons.append("标题样式缺失。")
    if 'toc1' not in normalized_styles:
        reasons.append("目录项样式缺失。")

    for token in BANNED_PLACEHOLDERS:
        if token in text:
            reasons.append(f"检测到历史占位词：{token}")

    if "<w:tbl" in document_xml:
        manual_review.append("输出中包含表格。")
    if "word/footnotes.xml" in parts:
        manual_review.append("输出中包含脚注或尾注。")

    if reasons:
        status = "NOT_SUPPORTED"
    elif manual_review:
        status = "MANUAL_REVIEW"
    else:
        status = "PASS"

    return {
        "status": status,
        "reasons": reasons,
        "manual_review": manual_review,
        "checked_file": str(path),
        "section_formats": section_formats,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Check SCNU thesis DOCX compliance.")
    parser.add_argument("docx_path", type=Path)
    args = parser.parse_args()
    report = build_report(args.docx_path)
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if report["status"] != "NOT_SUPPORTED" else 1


if __name__ == "__main__":
    raise SystemExit(main())
