from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable
from zipfile import ZipFile

from docx import Document

from pydantic import BaseModel, Field

from ..contracts import (
    AppendixSection,
    BodySection,
    CapabilityFlags,
    CoverFields,
    NormalizedThesis,
    ReferenceItem,
    SourceFeatures,
    SummarySection,
)
from ..errors import AppError

SPECIAL_TITLE_MAP = {
    "摘要": "abstract_cn",
    "中文摘要": "abstract_cn",
    "摘要摘要": "abstract_cn",
    "abstract": "abstract_en",
    "英文摘要": "abstract_en",
    "外文摘要": "abstract_en",
    "目录": "toc",
    "参考文献": "references",
    "references": "references",
    "致谢": "acknowledgements",
    "致謝": "acknowledgements",
    "acknowledgements": "acknowledgements",
    "acknowledgment": "acknowledgements",
    "附录": "appendices",
    "appendix": "appendices",
    "注释": "notes",
    "註釋": "notes",
    "注解": "notes",
    "notes": "notes",
}

COVER_FIELD_LABELS = {
    "title": ["论文题目", "论文名称", "题目", "毕业论文题目"],
    "advisor": ["指导老师", "指导教师"],
    "student_name": ["学生姓名", "姓名"],
    "student_id": ["学号"],
    "department": ["学院", "院系", "院（系）", "院(系)"],
    "major": ["专业"],
    "class_name": ["班级"],
    "graduation_time": ["毕业时间", "提交日期", "日期"],
}

MARKDOWN_HEADING = re.compile(r"^(#{1,4})\s+(.+)$")
CHAPTER_HEADING = re.compile(r"^第[一二三四五六七八九十百千0-9]+章")
NUMBERED_HEADING = re.compile(r"^(\d+(?:\.\d+){0,3})[\.、]?\s+(.+)$")
INLINE_FIELD_SPLIT = re.compile(r"[：:]")


@dataclass
class RawBlock:
    text: str
    style_name: str | None = None
    source_index: int = 0


@dataclass
class SectionDraft:
    kind: str
    title: str
    level: int
    lines: list[str] = field(default_factory=list)

    @property
    def content(self) -> str:
        return "\n".join(line.rstrip() for line in self.lines).strip()


def normalize_compact_text(value: str) -> str:
    return re.sub(r"[\s:：\-\._·•\(\)（）]+", "", (value or "")).strip().lower()


def split_keywords(text: str, english: bool) -> tuple[str, list[str]]:
    lines = [line.strip() for line in (text or "").splitlines()]
    body_lines: list[str] = []
    keywords_text = ""
    prefixes = ["keywords", "keyword"] if english else ["关键词", "關鍵詞"]
    for line in lines:
        normalized = line.lower().replace("：", ":")
        if any(normalized.startswith(prefix) for prefix in prefixes):
            keywords_text = line.split(":", 1)[-1].split("：", 1)[-1].strip()
            continue
        body_lines.append(line)
    keywords = [item.strip() for item in re.split(r"[，,；;、]", keywords_text) if item.strip()]
    return "\n".join(body_lines).strip(), keywords


def normalize_reference_text(text: str) -> tuple[str, str]:
    raw_text = " ".join((text or "").strip().split())
    normalized = raw_text
    normalized = re.sub(r"^[\[\(【]?\d+[\]\)】\.、\s]*", "", normalized)
    normalized = normalized.replace("：", ": ").replace("，", ", ")
    normalized = re.sub(r"\s+", " ", normalized).strip(" ;")
    normalized = re.sub(r"\s+([,.;:])", r"\1", normalized)
    normalized = re.sub(r"([A-Za-z])\.\s*([A-Za-z])", r"\1. \2", normalized)
    return raw_text, normalized


def build_capabilities(capabilities: CapabilityFlags) -> CapabilityFlags:
    return capabilities.model_copy(deep=True)


def make_source_features_docx(path: Path, document: Document) -> SourceFeatures:
    with ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
        footnotes_xml = archive.read("word/footnotes.xml").decode("utf-8", errors="ignore") if "word/footnotes.xml" in archive.namelist() else ""
    return SourceFeatures(
        table_count=len(document.tables),
        image_count=document_xml.count("<w:drawing"),
        footnote_count=footnotes_xml.count("<w:footnote ") if footnotes_xml else 0,
        textbox_count=document_xml.count("w:txbxContent"),
        shape_count=document_xml.count("<w:pict") + document_xml.count("<v:shape") + document_xml.count("<wps:wsp"),
        field_count=document_xml.count("<w:instrText"),
        rich_run_count=sum(1 for paragraph in document.paragraphs if len([run for run in paragraph.runs if run.text.strip()]) > 1),
    )


def build_source_features_text() -> SourceFeatures:
    return SourceFeatures()


def build_manual_review_flags(source_type: str, features: SourceFeatures) -> list[str]:
    flags: list[str] = []
    if source_type == "docx":
        if features.table_count:
            flags.append(f"检测到 {features.table_count} 个表格，导出后需人工复核。")
        if features.image_count:
            flags.append(f"检测到 {features.image_count} 个图片或绘图对象，导出后需人工复核。")
        if features.footnote_count:
            flags.append(f"检测到 {features.footnote_count} 个脚注或尾注引用，导出后需人工复核。")
        if features.textbox_count or features.shape_count:
            flags.append("检测到文本框或形状对象，导出后需人工复核。")
        if features.field_count:
            flags.append("检测到原始 Word 字段或域代码，导出后需人工复核。")
    return flags


def detect_heading(block: RawBlock) -> tuple[bool, str, int, str]:
    text = block.text.strip()
    if not text:
        return False, "", 0, "body"

    style = (block.style_name or "").strip().lower()
    if style.startswith("heading"):
        digits = re.sub(r"\D+", "", style)
        level = min(int(digits) if digits else 1, 4)
        title = text
        kind = SPECIAL_TITLE_MAP.get(normalize_compact_text(title), "body")
        return True, title, level, kind

    markdown_match = MARKDOWN_HEADING.match(text)
    if markdown_match:
        hashes, title = markdown_match.groups()
        kind = SPECIAL_TITLE_MAP.get(normalize_compact_text(title), "body")
        return True, title.strip(), len(hashes), kind

    normalized = normalize_compact_text(text)
    if normalized in SPECIAL_TITLE_MAP:
        return True, text, 1, SPECIAL_TITLE_MAP[normalized]

    if CHAPTER_HEADING.match(text):
        return True, text, 1, "body"

    numbered = NUMBERED_HEADING.match(text)
    if numbered:
        prefix, title = numbered.groups()
        level = prefix.count(".") + 1
        return True, text, min(level, 4), "body"

    return False, "", 0, "body"


def is_likely_cover_label(text: str) -> bool:
    normalized = normalize_compact_text(text)
    return any(normalized.startswith(normalize_compact_text(alias)) for aliases in COVER_FIELD_LABELS.values() for alias in aliases)


def is_likely_title(value: str) -> bool:
    text = value.strip()
    if not text:
        return False
    if len(text) > 80:
        return False
    normalized = normalize_compact_text(text)
    if normalized in SPECIAL_TITLE_MAP or normalized in {"正文", "引言", "绪论", "前言", "结论"}:
        return False
    if is_likely_cover_label(text):
        return False
    if CHAPTER_HEADING.match(text):
        return False
    if NUMBERED_HEADING.match(text):
        return False
    return True


def consume_following_lines(lines: list[str], start: int, *, allow_multiline: bool) -> tuple[str, set[int]]:
    consumed = {start}
    collected: list[str] = []
    for index in range(start + 1, len(lines)):
        candidate = lines[index].strip()
        if not candidate:
            if collected:
                break
            consumed.add(index)
            continue
        if is_likely_cover_label(candidate) or normalize_compact_text(candidate) in SPECIAL_TITLE_MAP:
            break
        collected.append(candidate)
        consumed.add(index)
        if not allow_multiline:
            break
    return "\n".join(collected).strip(), consumed


def extract_cover(front_lines: list[str]) -> tuple[CoverFields, set[int]]:
    cover = CoverFields()
    consumed: set[int] = set()

    if front_lines and front_lines[0].strip() == "华南师范大学":
        consumed.add(0)

    for index, line in enumerate(front_lines):
        stripped = line.strip()
        if not stripped or index in consumed:
            continue
        normalized = normalize_compact_text(stripped)
        for field, aliases in COVER_FIELD_LABELS.items():
            if not any(normalized.startswith(normalize_compact_text(alias)) for alias in aliases):
                continue
            _, _, remainder = stripped.partition("：")
            if not remainder:
                _, _, remainder = stripped.partition(":")
            value = remainder.strip()
            extra_consumed = {index}
            if not value:
                value, extra_consumed = consume_following_lines(front_lines, index, allow_multiline=field == "title")
            if field == "title" and value and not cover.title:
                cover.title = value
            elif field == "advisor" and value and not cover.advisor:
                cover.advisor = value
            elif field == "student_name" and value and not cover.student_name:
                cover.student_name = value
            elif field == "student_id" and value and not cover.student_id:
                cover.student_id = value
            elif field == "department" and value and not cover.department:
                cover.department = value
            elif field == "major" and value and not cover.major:
                cover.major = value
            elif field == "class_name" and value and not cover.class_name:
                cover.class_name = value
            elif field == "graduation_time" and value and not cover.graduation_time:
                cover.graduation_time = value
            consumed.update(extra_consumed)
            break

    if not cover.title:
        title_lines = [line.strip() for idx, line in enumerate(front_lines) if idx not in consumed and is_likely_title(line)]
        if title_lines:
            cover.title = "\n".join(title_lines[:2]).strip()
            for idx, line in enumerate(front_lines):
                if line.strip() in title_lines[:2]:
                    consumed.add(idx)

    return cover, consumed


def extract_raw_blocks_from_docx(path: Path) -> tuple[list[RawBlock], SourceFeatures]:
    try:
        document = Document(path)
    except Exception as exc:  # pragma: no cover
        raise AppError("PARSE_FAILED", "无法读取这个 .docx 的正文内容，请确认文件未损坏。", details={"reason": str(exc)}, status_code=400) from exc

    raw_blocks = [
        RawBlock(
            text=paragraph.text,
            style_name=paragraph.style.name if paragraph.style is not None else None,
            source_index=index,
        )
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip()
    ]
    return raw_blocks, make_source_features_docx(path, document)


def extract_raw_blocks_from_text(text: str) -> tuple[list[RawBlock], SourceFeatures]:
    raw_blocks = [RawBlock(text=line, style_name=None, source_index=index) for index, line in enumerate(text.splitlines()) if line.strip()]
    return raw_blocks, build_source_features_text()


def body_title_for_missing_input() -> str:
    return "正文"


def normalized_from_raw_blocks(
    raw_blocks: list[RawBlock],
    source_type: str,
    capabilities: CapabilityFlags,
    source_features: SourceFeatures,
) -> NormalizedThesis:
    if not raw_blocks:
        raise AppError("CONTENT_EMPTY", "文档中没有可用文本内容。", status_code=400)

    first_heading_index: int | None = None
    for index, block in enumerate(raw_blocks):
        is_heading, _, _, _ = detect_heading(block)
        if is_heading:
            first_heading_index = index
            break

    front_blocks = raw_blocks[: first_heading_index or 0]
    cover, consumed_front = extract_cover([block.text for block in front_blocks])
    front_remainder = [block.text.strip() for idx, block in enumerate(front_blocks) if idx not in consumed_front and block.text.strip()]

    drafts: list[SectionDraft] = []
    current: SectionDraft | None = None
    orphan_lines: list[str] = list(front_remainder)

    def flush_current() -> None:
        nonlocal current
        if current is None:
            return
        drafts.append(current)
        current = None

    for block in raw_blocks[first_heading_index or 0 :]:
        text = block.text.strip()
        if not text:
            continue
        is_heading, heading_title, level, kind = detect_heading(block)
        if is_heading:
            flush_current()
            current = SectionDraft(kind=kind, title=heading_title.strip(), level=level or 1)
            continue

        if current is None:
            orphan_lines.append(text)
            continue
        current.lines.append(text)

    flush_current()

    abstract_cn = SummarySection()
    abstract_en = SummarySection()
    body_sections: list[BodySection] = []
    references: list[ReferenceItem] = []
    appendices: list[AppendixSection] = []
    acknowledgements = ""
    notes = ""

    appendix_counter = 1
    for draft in drafts:
        content = draft.content
        if draft.kind == "abstract_cn" and not abstract_cn.content:
            body, keywords = split_keywords(content, english=False)
            abstract_cn = SummarySection(content=body, keywords=keywords)
            if not cover.title and orphan_lines and is_likely_title(orphan_lines[0]):
                cover.title = orphan_lines.pop(0)
        elif draft.kind == "abstract_en" and not abstract_en.content:
            body, keywords = split_keywords(content, english=True)
            abstract_en = SummarySection(content=body, keywords=keywords)
        elif draft.kind == "references":
            for line in content.splitlines():
                if not line.strip():
                    continue
                raw_text, normalized_text = normalize_reference_text(line)
                references.append(ReferenceItem(raw_text=raw_text, normalized_text=normalized_text))
        elif draft.kind == "appendices":
            title = draft.title.strip() or f"附录 {appendix_counter}"
            appendices.append(AppendixSection(id=f"appendix-{appendix_counter}", title=title, content=content))
            appendix_counter += 1
        elif draft.kind == "acknowledgements":
            acknowledgements = "\n".join(item for item in [acknowledgements, content] if item).strip()
        elif draft.kind == "notes":
            notes = "\n".join(item for item in [notes, content] if item).strip()
        elif draft.kind == "toc":
            continue
        else:
            body_sections.append(
                BodySection(
                    id=f"section-{len(body_sections) + 1}",
                    level=draft.level,
                    title=draft.title.strip() or body_title_for_missing_input(),
                    content=content,
                )
            )

    orphan_body = "\n".join(line for line in orphan_lines if line.strip()).strip()
    if orphan_body:
        if body_sections:
            body_sections[0].content = "\n".join(item for item in [orphan_body, body_sections[0].content] if item).strip()
        else:
            body_sections.append(
                BodySection(
                    id="section-1",
                    level=1,
                    title=body_title_for_missing_input(),
                    content=orphan_body,
                )
            )

    if not body_sections:
        raise AppError("CONTENT_EMPTY", "未识别到可映射为论文正文的内容。", status_code=400)

    missing_sections: list[str] = []
    for field in ["title", "advisor", "student_name", "student_id", "department", "major", "class_name", "graduation_time"]:
        if not getattr(cover, field).strip():
            missing_sections.append(f"cover.{field}")
    if not abstract_cn.content.strip():
        missing_sections.append("abstract_cn")
    if not abstract_en.content.strip():
        missing_sections.append("abstract_en")
    if not abstract_cn.keywords:
        missing_sections.append("keywords_cn")
    if abstract_en.content.strip() and not abstract_en.keywords:
        missing_sections.append("keywords_en")
    if not references:
        missing_sections.append("references")
    if not appendices or not any(item.content.strip() for item in appendices):
        missing_sections.append("appendices")
    if not acknowledgements.strip():
        missing_sections.append("acknowledgements")
    if not notes.strip():
        missing_sections.append("notes")

    manual_review_flags = build_manual_review_flags(source_type, source_features)
    warnings = manual_review_flags.copy()
    if "abstract_cn" in missing_sections:
        warnings.append("未识别到中文摘要，导出时会保留摘要章节留白。")
    if "abstract_en" in missing_sections:
        warnings.append("未识别到英文摘要，导出时会保留摘要章节留白。")
    if "references" in missing_sections:
        warnings.append("未识别到参考文献，导出时会保留参考文献章节留白。")
    if "appendices" in missing_sections:
        warnings.append("未识别到附录，导出时会保留附录章节留白。")
    if "acknowledgements" in missing_sections:
        warnings.append("未识别到致谢，导出时会保留致谢章节留白。")

    if not cover.title and body_sections:
        first_body_title = body_sections[0].title.strip()
        if is_likely_title(first_body_title) and len(re.sub(r"\s+", "", body_sections[0].content)) > 50:
            cover.title = first_body_title

    return NormalizedThesis(
        source_type="docx" if source_type == "docx" else "text",
        cover=cover,
        abstract_cn=abstract_cn,
        abstract_en=abstract_en,
        body_sections=body_sections,
        references=references,
        appendices=appendices,
        acknowledgements=acknowledgements,
        notes=notes,
        warnings=warnings,
        manual_review_flags=manual_review_flags,
        missing_sections=missing_sections,
        source_features=source_features,
        capabilities=build_capabilities(capabilities),
    )


def parse_docx_file(path: Path, capabilities: CapabilityFlags) -> NormalizedThesis:
    if path.suffix.lower() != ".docx":
        raise AppError("UNSUPPORTED_FILE_TYPE", "仅支持上传 .docx 文件。", status_code=400)
    raw_blocks, source_features = extract_raw_blocks_from_docx(path)
    return normalized_from_raw_blocks(raw_blocks, "docx", capabilities, source_features)


def normalize_text_input(text: str, capabilities: CapabilityFlags) -> NormalizedThesis:
    raw_blocks, source_features = extract_raw_blocks_from_text(text)
    return normalized_from_raw_blocks(raw_blocks, "text", capabilities, source_features)


# ─── Story2Paper Schema Mapper ─────────────────────────────────────────────────


class Story2PaperSchema(BaseModel):
    """Story2Paper 的 export_schema_json() 输出格式。"""
    paper_id: str | None = None
    title: str | None = None
    abstract_zh: str | None = ""
    abstract_en: str | None = ""
    sections: list[dict] = Field(default_factory=list)
    keywords: list[str] = Field(default_factory=list)
    references: list[str] = Field(default_factory=list)
    figures: list[dict] = Field(default_factory=list)
    tables: list[dict] = Field(default_factory=list)


def from_story2paper_json(
    raw: dict,
    cover_fields: CoverFields,
    capabilities: CapabilityFlags,
) -> NormalizedThesis:
    """
    将 Story2Paper 的 schema JSON 映射为 NormalizedThesis。

    封面字段（CoverFields）由用户在 UI 层手动填写后传入，
    其他字段从 Story2Paper pipeline 结果中提取。
    """
    schema = Story2PaperSchema.model_validate(raw)

    # 构建正文章节
    body_sections: list[BodySection] = []
    for idx, sec in enumerate(schema.sections):
        title = sec.get("title", f"第 {idx + 1} 节")
        content = sec.get("content", "")
        if not title.strip():
            title = body_title_for_missing_input()
        body_sections.append(
            BodySection(
                id=f"section-{idx + 1}",
                level=1,
                title=title.strip(),
                content=content.strip(),
            )
        )

    # 构建参考文献
    references: list[ReferenceItem] = []
    for ref_text in schema.references:
        raw_text = " ".join((ref_text or "").strip().split())
        references.append(
            ReferenceItem(
                raw_text=raw_text,
                normalized_text=raw_text,
                detected_type="",
            )
        )

    # Story2Paper 生成的论文缺少摘要/致谢/附录，手动标记为缺失
    missing_sections: list[str] = []
    if not schema.abstract_zh:
        missing_sections.append("abstract_cn")
    if not schema.abstract_en:
        missing_sections.append("abstract_en")
    if not references:
        missing_sections.append("references")
    missing_sections.append("appendices")
    missing_sections.append("acknowledgements")
    missing_sections.append("notes")

    # 警告信息
    warnings: list[str] = []
    if schema.figures:
        warnings.append(f"检测到 {len(schema.figures)} 个图表，导出后需人工复核。")
    if schema.tables:
        warnings.append(f"检测到 {len(schema.tables)} 个表格，导出后需人工复核。")
    warnings.append("AI 生成论文，导出后请仔细校对内容。")

    return NormalizedThesis(
        source_type="story2paper",
        cover=cover_fields,
        abstract_cn=SummarySection(content=(schema.abstract_zh or "").strip(), keywords=schema.keywords),
        abstract_en=SummarySection(content=(schema.abstract_en or "").strip(), keywords=[]),
        body_sections=body_sections,
        references=references,
        appendices=[],
        acknowledgements="",
        notes="",
        warnings=warnings,
        manual_review_flags=warnings.copy(),
        missing_sections=missing_sections,
        source_features=SourceFeatures(
            table_count=len(schema.tables),
            image_count=len(schema.figures),
        ),
        capabilities=build_capabilities(capabilities),
    )
