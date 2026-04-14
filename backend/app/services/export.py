from __future__ import annotations

import io
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from ..config import COVER_LOGO_PATH, DEBUG_OUTPUTS_DIR, DEBUG_PERSIST_ARTIFACTS, TEMPLATE_DOCX_PATH
from ..contracts import AppendixSection, BodySection, NormalizedThesis
from ..errors import AppError
from .precheck import run_precheck

CHINESE_BODY_FONT = "宋体"
CHINESE_HEADING_FONT = "黑体"
ENGLISH_FONT = "Times New Roman"
HEADER_FONT = "宋体"
FOOTER_FONT = "黑体"

HEADER_SUBTITLE_SEPARATOR_PATTERNS = (
    re.compile(r"^(?P<main>.+?)：\s*(?P<sub>.+)$"),
    re.compile(r"^(?P<main>.+?):\s+(?P<sub>.+)$"),
    re.compile(r"^(?P<main>.+?)(?:——|--|—)\s*(?P<sub>.+)$"),
    re.compile(r"^(?P<main>.+?)\s+-\s+(?P<sub>.+)$"),
    re.compile(r"^(?P<main>.+?)\s*\|\s*(?P<sub>.+)$"),
)
HEADER_PARENTHESES_SUBTITLE_PATTERN = re.compile(r"^(?P<main>.+?)[（(]\s*(?P<sub>[^()（）]{1,40})\s*[）)]$")


@dataclass
class RenderedBodySection:
    title: str
    level: int
    content: str


@dataclass
class RenderPlan:
    header_title: str
    body_sections: list[RenderedBodySection]


def normalize_text_block(text: str) -> str:
    return "\n".join(line.rstrip() for line in (text or "").strip().splitlines()).strip()


def primary_title_line(title: str) -> str:
    normalized = normalize_text_block(title)
    if not normalized:
        return ""
    for line in normalized.splitlines():
        stripped = line.strip()
        if stripped:
            return stripped
    return ""


def has_title_letters(text: str) -> bool:
    return bool(re.search(r"[A-Za-z\u4e00-\u9fff]", text or ""))


def looks_like_version_or_year(text: str) -> bool:
    candidate = (text or "").strip()
    return bool(re.fullmatch(r"\d{2,4}(?:[./-]\d{1,2}){0,2}(?:版)?", candidate))


def clean_header_main_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip()).rstrip("：:|-—– ")


def can_strip_subtitle(main_title: str, subtitle: str) -> bool:
    main = clean_header_main_text(main_title)
    sub = re.sub(r"\s+", " ", (subtitle or "").strip())
    if not main or not sub:
        return False
    if not has_title_letters(main) or not has_title_letters(sub):
        return False
    if len(sub) > 40:
        return False
    if re.fullmatch(r"[A-Z]{1,6}", sub):
        return False
    if looks_like_version_or_year(sub):
        return False
    return True


def strip_subtitle_for_header(title: str) -> str:
    line = primary_title_line(title)
    for pattern in HEADER_SUBTITLE_SEPARATOR_PATTERNS:
        match = pattern.match(line)
        if not match:
            continue
        main = match.group("main")
        subtitle = match.group("sub")
        if can_strip_subtitle(main, subtitle):
            return clean_header_main_text(main)

    parenthetical = HEADER_PARENTHESES_SUBTITLE_PATTERN.match(line)
    if parenthetical and can_strip_subtitle(parenthetical.group("main"), parenthetical.group("sub")):
        return clean_header_main_text(parenthetical.group("main"))
    return clean_header_main_text(line)


def persist_debug_copy(label: str, payload: bytes, suffix: str) -> None:
    if not DEBUG_PERSIST_ARTIFACTS:
        return
    DEBUG_OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    target = DEBUG_OUTPUTS_DIR / f"{timestamp}-{label}.{suffix}"
    target.write_bytes(payload)


def clear_document(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def ensure_settings_update_fields(document: Document) -> None:
    settings = document.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def set_rfonts(target, east_asia: str, ascii_name: str | None = None) -> None:
    ascii_font = ascii_name or east_asia
    rpr = target.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)


def set_style_font(style, east_asia: str, size_pt: int, *, ascii_name: str | None = None, bold: bool = False) -> None:
    style.font.name = ascii_name or east_asia
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    set_rfonts(style._element, east_asia, ascii_name)


def set_run_font(run, east_asia: str, size_pt: int, *, ascii_name: str | None = None, bold: bool = False, underline: bool = False) -> None:
    ascii_font = ascii_name or east_asia
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.underline = underline
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)


def ensure_style(document: Document, style_name: str) -> None:
    styles = document.styles
    try:
        styles[style_name]
    except KeyError:
        styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)


def configure_styles(document: Document) -> None:
    styles = document.styles

    set_style_font(styles["Normal"], CHINESE_BODY_FONT, 12, ascii_name=ENGLISH_FONT)
    styles["Normal"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    styles["Normal"].paragraph_format.line_spacing = 1.25
    styles["Normal"].paragraph_format.first_line_indent = Cm(0.74)
    styles["Normal"].paragraph_format.space_before = Pt(0)
    styles["Normal"].paragraph_format.space_after = Pt(0)

    set_style_font(styles["Title"], CHINESE_HEADING_FONT, 18, bold=True)

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        ensure_style(document, style_name)
    set_style_font(styles["Heading 1"], CHINESE_HEADING_FONT, 18, bold=True)
    set_style_font(styles["Heading 2"], CHINESE_HEADING_FONT, 14, bold=True)
    set_style_font(styles["Heading 3"], CHINESE_HEADING_FONT, 12, bold=True)
    set_style_font(styles["Heading 4"], CHINESE_HEADING_FONT, 12, bold=True)
    styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for style_name in ["Heading 2", "Heading 3", "Heading 4"]:
        styles[style_name].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for style_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        styles[style_name].paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        styles[style_name].paragraph_format.line_spacing = 1.25
        styles[style_name].paragraph_format.space_before = Pt(0)
        styles[style_name].paragraph_format.space_after = Pt(0)

    for toc_name in ["TOC 1", "TOC 2", "TOC 3", "TOC 4"]:
        try:
            style = styles[toc_name]
        except KeyError:
            continue
        set_style_font(style, CHINESE_BODY_FONT, 12, ascii_name=ENGLISH_FONT)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)


def set_section_geometry(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)
    pg_mar = section._sectPr.find(qn("w:pgMar"))
    if pg_mar is not None:
        pg_mar.set(qn("w:gutter"), str(int(Cm(0.5).twips)))


def set_section_page_number_format(section, fmt: str | None, start: int | None = None) -> None:
    sect_pr = section._sectPr
    page_number = sect_pr.find(qn("w:pgNumType"))
    if fmt is None and start is None:
        if page_number is not None:
            sect_pr.remove(page_number)
        return
    if page_number is None:
        page_number = OxmlElement("w:pgNumType")
        sect_pr.append(page_number)
    if fmt is not None:
        page_number.set(qn("w:fmt"), fmt)
    if start is not None:
        page_number.set(qn("w:start"), str(start))


def clear_paragraph(paragraph) -> None:
    element = paragraph._element
    for child in list(element):
        element.remove(child)


def add_field(run, instruction: str) -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr_text, fld_separate, fld_end])


def configure_header(section, header_title: str) -> None:
    section.header.is_linked_to_previous = False
    header = section.header
    paragraph = header.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(header_title)
    set_run_font(run, HEADER_FONT, 10, ascii_name=ENGLISH_FONT)


def configure_footer(section) -> None:
    section.footer.is_linked_to_previous = False
    footer = section.footer
    paragraph = footer.paragraphs[0]
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    set_run_font(run, FOOTER_FONT, 10, ascii_name=FOOTER_FONT, bold=True)
    add_field(run, "PAGE")


def blank_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_paragraph(section.header.paragraphs[0])
    clear_paragraph(section.footer.paragraphs[0])


def weighted_title_length(value: str) -> float:
    total = 0.0
    for char in value:
        total += 0.5 if ord(char) < 128 else 1.0
    return total


def extract_header_title(title: str, *, max_length: int = 28) -> str:
    main_line = strip_subtitle_for_header(title)
    if not main_line:
        return ""
    if weighted_title_length(main_line) <= max_length:
        return main_line
    result = ""
    current = 0.0
    for char in main_line:
        width = 0.5 if ord(char) < 128 else 1.0
        if current + width > max_length:
            break
        result += char
        current += width
    return result.strip()


def truncate_header_title(title: str) -> str:
    return extract_header_title(title, max_length=28)


def strip_existing_body_prefix(title: str) -> str:
    value = (title or "").strip()
    value = re.sub(r"^第[一二三四五六七八九十百千0-9]+章\s*", "", value)
    value = re.sub(r"^\d+(?:\.\d+){0,3}[\.、]?\s*", "", value)
    return value.strip() or "正文"


def section_number(section: BodySection, counters: list[int]) -> str:
    level = min(max(section.level, 1), 4)
    for index in range(level - 1):
        if counters[index] == 0:
            counters[index] = 1
    counters[level - 1] += 1
    for index in range(level, len(counters)):
        counters[index] = 0
    return ".".join(str(counters[index]) for index in range(level)) + "."


def build_render_plan(thesis: NormalizedThesis) -> RenderPlan:
    header_title = truncate_header_title(thesis.cover.title)
    counters = [0, 0, 0, 0]
    rendered_sections: list[RenderedBodySection] = []

    for section in thesis.body_sections:
        level = min(max(section.level, 1), 4)
        numbering = section_number(section, counters).rstrip(".")
        clean_title = strip_existing_body_prefix(section.title)
        display_title = f"{numbering} {clean_title}".strip()
        rendered_sections.append(RenderedBodySection(title=display_title, level=level, content=section.content))

    return RenderPlan(header_title=header_title, body_sections=rendered_sections)


def load_template_document() -> Document:
    if not TEMPLATE_DOCX_PATH.exists():
        raise AppError(
            "TEMPLATE_DEPENDENCY_MISSING",
            "Word 模板不存在，无法生成导出结果。",
            details={"template": str(TEMPLATE_DOCX_PATH)},
            status_code=500,
        )
    try:
        document = Document(TEMPLATE_DOCX_PATH)
    except Exception as exc:  # pragma: no cover
        raise AppError(
            "TEMPLATE_DEPENDENCY_MISSING",
            "Word 模板无法加载，请检查模板文件是否损坏。",
            details={"template": str(TEMPLATE_DOCX_PATH), "reason": str(exc)},
            status_code=500,
        ) from exc
    clear_document(document)
    ensure_settings_update_fields(document)
    configure_styles(document)
    set_section_geometry(document.sections[0])
    return document


def add_blank_body_paragraph(document: Document, *, english: bool = False, hanging: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Cm(0) if hanging else Cm(0.74)
    if hanging:
        paragraph.paragraph_format.left_indent = Cm(0.74)
        paragraph.paragraph_format.first_line_indent = Cm(-0.74)
    run = paragraph.add_run("")
    if english:
        set_run_font(run, ENGLISH_FONT, 12, ascii_name=ENGLISH_FONT)
    else:
        set_run_font(run, CHINESE_BODY_FONT, 12, ascii_name=ENGLISH_FONT)


def add_body_paragraphs(document: Document, text: str, *, english: bool = False, hanging: bool = False) -> None:
    chunks = [chunk.strip() for chunk in normalize_text_block(text).split("\n\n") if chunk.strip()]
    if not chunks:
        add_blank_body_paragraph(document, english=english, hanging=hanging)
        return
    for chunk in chunks:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        if hanging:
            paragraph.paragraph_format.left_indent = Cm(0.74)
            paragraph.paragraph_format.first_line_indent = Cm(-0.74)
        else:
            paragraph.paragraph_format.first_line_indent = Cm(0.74)
        run = paragraph.add_run(chunk)
        if english:
            set_run_font(run, ENGLISH_FONT, 12, ascii_name=ENGLISH_FONT)
        else:
            set_run_font(run, CHINESE_BODY_FONT, 12, ascii_name=ENGLISH_FONT)


def add_heading(document: Document, text: str, level: int, *, center: bool | None = None) -> None:
    style_name = f"Heading {min(max(level, 1), 4)}"
    paragraph = document.add_paragraph(style=style_name)
    if center is not None:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, CHINESE_HEADING_FONT, 18, bold=True)
    elif level == 2:
        set_run_font(run, CHINESE_HEADING_FONT, 14, bold=True)
    else:
        set_run_font(run, CHINESE_HEADING_FONT, 12, bold=True)


def add_page_break(document: Document) -> None:
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_logo(document: Document) -> None:
    if not COVER_LOGO_PATH.exists():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(COVER_LOGO_PATH), width=Inches(1.0))


def cover_blank(value: str, width: int) -> str:
    if value.strip():
        return value.strip()
    return "_" * width


def add_cover_field_line(document: Document, label: str, value: str, *, width: int = 28) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.left_indent = Cm(3.8)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    label_run = paragraph.add_run(f"{label}：")
    set_run_font(label_run, CHINESE_BODY_FONT, 12, ascii_name=ENGLISH_FONT)
    value_run = paragraph.add_run(cover_blank(value, width))
    set_run_font(value_run, CHINESE_BODY_FONT, 12, ascii_name=ENGLISH_FONT)


def render_cover(document: Document, thesis: NormalizedThesis) -> None:
    cover = thesis.cover
    blank_header_footer(document.sections[0])
    set_section_page_number_format(document.sections[0], None, None)

    for _ in range(2):
        document.add_paragraph()
    add_logo(document)

    school_paragraph = document.add_paragraph()
    school_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    school_run = school_paragraph.add_run(cover.school.strip() or "华南师范大学")
    set_run_font(school_run, CHINESE_HEADING_FONT, 22, bold=True)

    doc_type_paragraph = document.add_paragraph()
    doc_type_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_type_run = doc_type_paragraph.add_run("本科毕业论文")
    set_run_font(doc_type_run, CHINESE_HEADING_FONT, 18, bold=True)

    for _ in range(4):
        document.add_paragraph()

    title_text = cover.title.strip().splitlines() if cover.title.strip() else []
    if title_text:
        for line in title_text[:2]:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line.strip())
            set_run_font(run, CHINESE_HEADING_FONT, 18, bold=True)
    else:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run("_" * 28)
        set_run_font(run, CHINESE_HEADING_FONT, 18, bold=True)

    for _ in range(3):
        document.add_paragraph()

    add_cover_field_line(document, "论文题目", " ".join(title_text))
    add_cover_field_line(document, "指导教师", cover.advisor, width=16)
    add_cover_field_line(document, "学生姓名", cover.student_name, width=16)
    add_cover_field_line(document, "学号", cover.student_id, width=16)
    add_cover_field_line(document, "学院", cover.department, width=18)
    add_cover_field_line(document, "专业", cover.major, width=18)
    add_cover_field_line(document, "班级", cover.class_name, width=18)
    add_cover_field_line(document, "毕业时间", cover.graduation_time, width=16)


def render_summary_section(document: Document, heading: str, content: str, keywords_label: str, keywords: list[str], *, english: bool) -> None:
    add_heading(document, heading, 1, center=True)
    add_body_paragraphs(document, content, english=english)
    keyword_paragraph = document.add_paragraph()
    keyword_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    keyword_paragraph.paragraph_format.line_spacing = 1.25
    keyword_paragraph.paragraph_format.space_before = Pt(0)
    keyword_paragraph.paragraph_format.space_after = Pt(0)
    keyword_paragraph.paragraph_format.first_line_indent = Cm(0.74)
    label_run = keyword_paragraph.add_run(f"{keywords_label}：")
    if english:
        set_run_font(label_run, ENGLISH_FONT, 12, ascii_name=ENGLISH_FONT, bold=True)
        value_run = keyword_paragraph.add_run(", ".join(keywords))
        set_run_font(value_run, ENGLISH_FONT, 12, ascii_name=ENGLISH_FONT)
    else:
        set_run_font(label_run, CHINESE_BODY_FONT, 12, ascii_name=ENGLISH_FONT, bold=True)
        value_run = keyword_paragraph.add_run("，".join(keywords))
        set_run_font(value_run, CHINESE_BODY_FONT, 12, ascii_name=ENGLISH_FONT)
    if not keywords:
        blank_run = keyword_paragraph.add_run("_" * 18)
        if english:
            set_run_font(blank_run, ENGLISH_FONT, 12, ascii_name=ENGLISH_FONT)
        else:
            set_run_font(blank_run, CHINESE_BODY_FONT, 12, ascii_name=ENGLISH_FONT)


def render_toc(document: Document) -> None:
    add_heading(document, "目  录", 1, center=True)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    add_field(run, 'TOC \\o "1-4" \\h \\z \\u')


def render_body(document: Document, render_plan: RenderPlan) -> None:
    for index, section in enumerate(render_plan.body_sections):
        add_heading(document, section.title, 1 if section.level == 1 else section.level, center=section.level == 1)
        add_body_paragraphs(document, section.content)
        if index != len(render_plan.body_sections) - 1:
            continue


def render_reference_section(document: Document, references) -> None:
    add_page_break(document)
    add_heading(document, "参考文献", 1, center=True)
    if not references:
        add_blank_body_paragraph(document, hanging=True)
        return
    for index, item in enumerate(references, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.left_indent = Cm(0.74)
        paragraph.paragraph_format.first_line_indent = Cm(-0.74)
        run = paragraph.add_run(f"[{index}] {item.normalized_text or item.raw_text}")
        set_run_font(run, CHINESE_BODY_FONT, 12, ascii_name=ENGLISH_FONT)


def render_optional_text_section(document: Document, heading: str, text: str) -> None:
    add_page_break(document)
    add_heading(document, heading, 1, center=True)
    add_body_paragraphs(document, text)


def render_appendices(document: Document, appendices: list[AppendixSection]) -> None:
    add_page_break(document)
    add_heading(document, "附录", 1, center=True)
    if not appendices:
        add_blank_body_paragraph(document)
        return
    for appendix in appendices:
        add_heading(document, appendix.title.strip() or "附录", 2, center=False)
        add_body_paragraphs(document, appendix.content)


def validate_for_export(thesis: NormalizedThesis) -> None:
    precheck = run_precheck(thesis)
    if precheck.summary.blocking_count > 0:
        raise AppError(
            "FIELD_MISSING",
            "预检仍存在阻塞项，无法导出 Word 文件。",
            details={
                "blocking_count": precheck.summary.blocking_count,
                "issues": [item.model_dump() for item in precheck.issues if item.severity == "blocking"],
            },
            status_code=400,
        )


def export_docx(thesis: NormalizedThesis) -> bytes:
    validate_for_export(thesis)
    render_plan = build_render_plan(thesis)

    try:
        document = load_template_document()

        render_cover(document, thesis)

        front_section = document.add_section(WD_SECTION.NEW_PAGE)
        set_section_geometry(front_section)
        set_section_page_number_format(front_section, "upperRoman", 1)
        configure_header(front_section, render_plan.header_title)
        configure_footer(front_section)

        render_summary_section(document, "摘  要", thesis.abstract_cn.content, "关键词", thesis.abstract_cn.keywords, english=False)
        add_page_break(document)
        render_summary_section(document, "Abstract", thesis.abstract_en.content, "Keywords", thesis.abstract_en.keywords, english=True)
        add_page_break(document)
        render_toc(document)

        body_section = document.add_section(WD_SECTION.NEW_PAGE)
        set_section_geometry(body_section)
        set_section_page_number_format(body_section, "decimal", 1)
        configure_header(body_section, render_plan.header_title)
        configure_footer(body_section)

        render_body(document, render_plan)
        if thesis.notes.strip():
            render_optional_text_section(document, "注释", thesis.notes)
        render_reference_section(document, thesis.references)
        render_appendices(document, thesis.appendices)
        render_optional_text_section(document, "致谢", thesis.acknowledgements)

        buffer = io.BytesIO()
        document.save(buffer)
        payload = buffer.getvalue()
        persist_debug_copy("word-thesis", payload, "docx")
        return payload
    except AppError:
        raise
    except Exception as exc:  # pragma: no cover
        raise AppError("EXPORT_FAILED", "导出 Word 文件失败，请稍后重试。", details={"reason": str(exc)}, status_code=500) from exc
