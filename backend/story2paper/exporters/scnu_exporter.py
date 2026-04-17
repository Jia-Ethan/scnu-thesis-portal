"""
scnu_exporter — Story2Paper draft → scnu-thesis-portal 格式导出

支持两种输出格式：
1. schema JSON  — 对齐 scnu-thesis-portal 的 NormalizedSchema，可直接导入 portal 做格式校验
2. .docx        — 符合华南师范大学论文格式的 Word 文档，可直接提交或导入 portal
"""

from __future__ import annotations
import json
import os
from pathlib import Path
from typing import Any

# Optional: python-docx for .docx export
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ─── Schema JSON Exporter ──────────────────────────────────────────────────────

def export_schema_json(draft_output: dict, output_dir: str = "output") -> str:
    """
    将 Story2Paper 的 pipeline 输出转换为 scnu-thesis-portal 可消费的 NormalizedSchema JSON.

    Args:
        draft_output: {
            "paper_id": str,
            "outline": {...},
            "section_drafts": [{"section_id", "title", "content"}],
            "contract": {...}
        }
        output_dir: 输出目录

    Returns:
        输出文件的绝对路径
    """
    os.makedirs(output_dir, exist_ok=True)
    paper_id = draft_output.get("paper_id", "unknown")
    outline = draft_output.get("outline", {})
    sections = draft_output.get("section_drafts", [])
    contract = draft_output.get("contract", {})

    # Map section drafts to schema format
    schema_sections = []
    for draft in sections:
        schema_sections.append({
            "section_id": draft.get("section_id", ""),
            "title": draft.get("title", ""),
            "content": draft.get("content", ""),
        })

    # Map figures from contract
    schema_figures = [
        {
            "id": fig.get("figure_id", f"fig_{i}"),
            "caption": fig.get("caption", ""),
            "position": fig.get("position", ""),
            "semantic_commitment": fig.get("semantic_commitment", ""),
        }
        for i, fig in enumerate(contract.get("figures", []))
    ]

    # Map tables from contract
    schema_tables = [
        {
            "id": tbl.get("table_id", f"tbl_{i}"),
            "caption": tbl.get("caption", ""),
            "columns": tbl.get("columns", []),
            "position": tbl.get("position", ""),
            "data_semantics": tbl.get("data_semantics", ""),
        }
        for i, tbl in enumerate(contract.get("tables", []))
    ]

    schema = {
        "$schema": "https://github.com/Jia-Ethan/scnu-thesis-portal/blob/main/normalized-schema.json",
        "paper_id": paper_id,
        "title": outline.get("title", "Untitled"),
        "abstract_zh": _extract_abstract(sections, "zh"),
        "abstract_en": _extract_abstract(sections, "en"),
        "sections": schema_sections,
        "figures": schema_figures,
        "tables": schema_tables,
        "novelty_statement": outline.get("novelty_statement", ""),
        "keywords": _extract_keywords(sections),
        "references": [c.get("full_reference", "") for c in contract.get("citations", [])],
        "source": "Story2Paper",
        "exported_at": _iso_now(),
    }

    out_path = os.path.join(output_dir, f"{paper_id}_schema.json")
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(schema, f, ensure_ascii=False, indent=2)
    return out_path


# ─── DOCX Exporter ────────────────────────────────────────────────────────────

def export_docx(draft_output: dict, output_dir: str = "output") -> str:
    """
    将 Story2Paper 的 pipeline 输出导出为符合华师论文格式的 .docx 文档。

    华师论文格式要求：
    - 中文宋体标题，英文 Times New Roman
    - 正文宋体小四号（12pt），英文 Times New Roman
    - 1.5 倍行距
    - 页面边距：上下 2.54cm，左右 3cm

    Args:
        draft_output: 同上
        output_dir: 输出目录

    Returns:
        输出文件的绝对路径
    """
    if not HAS_DOCX:
        raise ImportError(
            "python-docx not installed. Run: pip install python-docx\n"
            "Or use export_schema_json() instead."
        )

    os.makedirs(output_dir, exist_ok=True)
    paper_id = draft_output.get("paper_id", "unknown")
    outline = draft_output.get("outline", {})
    sections = draft_output.get("section_drafts", [])

    doc = Document()

    # Set page margins (华师标准: 上下 2.54cm, 左右 3cm)
    section = doc.sections[0]
    section.top_margin = Inches(1.0)    # 2.54cm ≈ 1 inch
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.18)    # 3cm ≈ 1.18 inch
    section.right_margin = Inches(1.18)

    # Title
    title = outline.get("title", "Untitled")
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_font(heading.runs[0], size=22, bold=True, zh=True)

    # Abstract (if found)
    abstract_zh = _extract_abstract(sections, "zh")
    if abstract_zh:
        doc.add_paragraph()
        abstract_para = doc.add_paragraph()
        abstract_para.add_run("摘要").bold = True
        abstract_para.add_run(f"\n{abstract_zh}")
        abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    abstract_en = _extract_abstract(sections, "en")
    if abstract_en:
        abstract_para = doc.add_paragraph()
        abstract_para.add_run("Abstract").bold = True
        abstract_para.add_run(f"\n{abstract_en}")
        abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Keywords
    keywords = _extract_keywords(sections)
    if keywords:
        kw_para = doc.add_paragraph()
        kw_para.add_run("关键词：").bold = True
        kw_para.add_run("；".join(keywords))
        kw_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()

    # Sections
    for draft in sections:
        sec_title = draft.get("title", "")
        sec_content = draft.get("content", "")

        # Skip abstract sections (already handled)
        if any(k in sec_title.lower() for k in ["abstract", "摘要"]):
            continue

        # Section heading (华师格式: 一级标题黑体三号)
        heading = doc.add_heading(sec_title, level=1)
        for run in heading.runs:
            _apply_font(run, size=16, bold=True, zh=True)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Body paragraphs
        for para_text in sec_content.split("\n\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            if para_text.startswith("#"):
                continue  # skip markdown headers
            para = doc.add_paragraph(para_text)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in para.runs:
                _apply_font(run, size=12, bold=False, zh=True)
            para.paragraph_format.line_spacing = 1.5

    out_path = os.path.join(output_dir, f"{paper_id}_paper.docx")
    doc.save(out_path)
    return out_path


# ─── Helpers ───────────────────────────────────────────────────────────────────

def _extract_abstract(sections: list, lang: str) -> str:
    """Extract abstract for given language from section drafts."""
    for draft in sections:
        title = draft.get("title", "").lower()
        content = draft.get("content", "")
        if (lang == "zh" and "摘要" in title) or (lang == "en" and "abstract" in title):
            return content.strip()
    return ""


def _extract_keywords(sections: list) -> list[str]:
    """Extract keywords from section drafts."""
    for draft in sections:
        title = draft.get("title", "")
        content = draft.get("content", "")
        if "关键词" in title or "keywords" in title.lower():
            # Split by common separators
            for sep in ["，", ",", "；", ";"]:
                if sep in content:
                    return [k.strip() for k in content.split(sep) if k.strip()]
            return [content.strip()]
    return []


def _apply_font(run, size: int, bold: bool, zh: bool):
    """Apply Chinese/academic font styling to a run."""
    run.font.size = Pt(size)
    run.font.bold = bold
    if zh:
        run.font.name = "宋体"    # SimSun for Chinese body
        run._element.rPr.rFonts.set(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}eastAsia",
            "宋体"
        )
    else:
        run.font.name = "Times New Roman"


def _iso_now() -> str:
    from datetime import datetime
    return datetime.now().isoformat()


# ─── Convenience ──────────────────────────────────────────────────────────────

def export_all(draft_output: dict, output_dir: str = "output") -> dict:
    """
    同时导出 JSON schema 和 DOCX（如果依赖可用）。
    Returns paths to both output files.
    """
    paths = {}
    try:
        paths["schema_json"] = export_schema_json(draft_output, output_dir)
    except Exception as e:
        paths["schema_json_error"] = str(e)

    if HAS_DOCX:
        try:
            paths["docx"] = export_docx(draft_output, output_dir)
        except Exception as e:
            paths["docx_error"] = str(e)
    else:
        paths["docx_skipped"] = "python-docx not installed"

    return paths
